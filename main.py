import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document

def wczytaj_plik():
    filepath = filedialog.askopenfilename(filetypes=[("Pliki tekstowe", "*.txt *.docx")])
    if not filepath:
        return

    tekst = ""
    if filepath.endswith(".txt"):
        with open(filepath, "r", encoding="utf-8") as f:
            tekst = f.read()
    elif filepath.endswith(".docx"):
        doc = Document(filepath)
        tekst = "\n".join([p.text for p in doc.paragraphs])
    else:
        messagebox.showerror("Błąd", "Nieobsługiwany format pliku.")
        return

    tekst_input.delete("1.0", tk.END)
    tekst_input.insert(tk.END, tekst)

def wczytaj_slowa():
    filepath = filedialog.askopenfilename(filetypes=[("Pliki tekstowe", "*.txt")])
    if not filepath:
        return

    with open(filepath, "r", encoding="utf-8") as f:
        slowa = f.read().replace("\n", ",")
    slowa_input.delete(0, tk.END)
    slowa_input.insert(0, slowa)

def cenzuruj_tekst():
    tekst = tekst_input.get("1.0", tk.END)
    slowa_do_cenzury = slowa_input.get()
    slowa = [s.strip().lower() for s in slowa_do_cenzury.split(",")]

    for slowo in slowa:
        if slowo:
            tekst = tekst.replace(slowo, "*" * len(slowo))
            tekst = tekst.replace(slowo.capitalize(), "*" * len(slowo))
            tekst = tekst.replace(slowo.upper(), "*" * len(slowo))
    
    wynik_output.delete("1.0", tk.END)
    wynik_output.insert(tk.END, tekst)

def zapisz_do_pliku():
    tekst = wynik_output.get("1.0", tk.END).strip()
    if not tekst:
        messagebox.showwarning("Brak danych", "Brak tekstu do zapisania.")
        return

    filepath = filedialog.asksaveasfilename(defaultextension=".txt",
                                             filetypes=[
                                                 ("Plik tekstowy", "*.txt"),
                                                 ("Plik Word (.docx)", "*.docx")
                                             ],
                                             title="Zapisz ocenzurowany tekst jako")
    if not filepath:
        return

    try:
        if filepath.endswith(".txt"):
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(tekst)
        elif filepath.endswith(".docx"):
            doc = Document()
            for line in tekst.split("\n"):
                doc.add_paragraph(line)
            doc.save(filepath)
        else:
            messagebox.showerror("Błąd", "Nieobsługiwany format pliku.")
            return

        messagebox.showinfo("Sukces", f"Zapisano do pliku:\n{filepath}")
    except Exception as e:
        messagebox.showerror("Błąd zapisu", f"Wystąpił błąd: {e}")

# GUI
root = tk.Tk()
root.title("Cenzor Tekstu z Pliku")
root.geometry("700x700")

tk.Button(root, text="Wczytaj tekst z pliku (.txt / .docx)", command=wczytaj_plik).pack(pady=5)

tekst_input = tk.Text(root, height=10, width=80)
tekst_input.pack()

frame = tk.Frame(root)
frame.pack()
tk.Label(frame, text="Słowa do cenzury (oddzielone przecinkami):").pack(side=tk.LEFT)
slowa_input = tk.Entry(frame, width=50)
slowa_input.pack(side=tk.LEFT, padx=5)
tk.Button(frame, text="Wczytaj z pliku", command=wczytaj_slowa).pack(side=tk.LEFT)

tk.Button(root, text="Cenzuruj", command=cenzuruj_tekst).pack(pady=10)

tk.Label(root, text="Ocenzurowany tekst:").pack()
wynik_output = tk.Text(root, height=10, width=80)
wynik_output.pack()

tk.Button(root, text="Zapisz do pliku (.txt / .docx)", command=zapisz_do_pliku).pack(pady=10)

root.mainloop()
